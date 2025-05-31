from pathlib import Path
from typing import Iterator, Dict, Any, Optional, List
from datetime import datetime
import docx
import re
import logging
from docx.opc.exceptions import PackageNotFoundError

from .base_loader import BaseLoader
from dataset.scripts.converters import _calculate_sha256

logger = logging.getLogger(__name__)

class DocxLoader(BaseLoader):
    """Loader para archivos DOCX (Microsoft Word)."""
    
    def __init__(self, file_path: str | Path, tipo: str = 'escritos', encoding: str = 'utf-8'):
        """
        Inicializa el loader de DOCX.
        
        Args:
            file_path: Ruta al archivo DOCX
            tipo: Tipo de contenido ('escritos', 'poemas', 'canciones')
            encoding: Codificación para leer el texto (por defecto utf-8)
        """
        super().__init__(file_path)
        self.tipo = tipo.lower()
        self.encoding = encoding
        logger.debug(f"Inicializando DocxLoader para archivo: {file_path}")
        
    def _extract_date_from_core_properties(self, doc) -> Optional[str]:
        """Intenta extraer la fecha de las propiedades del documento."""
        try:
            # Intenta obtener la fecha de creación
            if doc.core_properties.created:
                return doc.core_properties.created.strftime('%Y-%m-%d')
            # Si no hay fecha de creación, intenta la de modificación
            if doc.core_properties.modified:
                return doc.core_properties.modified.strftime('%Y-%m-%d')
        except Exception as e:
            logger.warning(f"Error al extraer fecha de propiedades: {str(e)}")
        return None
    
    def _extract_date_from_filename(self) -> Optional[str]:
        """Intenta extraer una fecha del nombre del archivo."""
        # Patrones comunes de fecha en nombres de archivo
        patterns = [
            r'(\d{4})[_-](\d{2})[_-](\d{2})',  # 2024-01-30, 2024_01_30
            r'(\d{2})[_-](\d{2})[_-](\d{4})',  # 30-01-2024, 30_01_2024
            r'(\d{4})(\d{2})(\d{2})'           # 20240130
        ]
        
        filename = self.file_path.stem
        
        for pattern in patterns:
            match = re.search(pattern, filename)
            if match:
                try:
                    if len(match.group(1)) == 4:  # Año primero
                        year, month, day = match.groups()
                    else:  # Día primero
                        day, month, year = match.groups()
                    
                    # Validar fecha
                    date = datetime(int(year), int(month), int(day))
                    return date.strftime('%Y-%m-%d')
                except ValueError:
                    continue
        
        return None
    
    def _clean_text(self, text: str) -> str:
        """Limpia el texto de caracteres no deseados y espacios extra."""
        try:
            # Intenta decodificar si es bytes
            if isinstance(text, bytes):
                text = text.decode(self.encoding, errors='replace')
            
            # Elimina caracteres de control excepto saltos de línea
            text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]', '', text)
            # Normaliza saltos de línea
            text = text.replace('\r\n', '\n').replace('\r', '\n')
            # Elimina espacios múltiples y al inicio/final
            text = re.sub(r' +', ' ', text).strip()
            
            return text
        except Exception as e:
            logger.warning(f"Error al limpiar texto: {str(e)}")
            # Retorna una versión segura del texto
            return str(text).encode('ascii', errors='replace').decode('ascii')
    
    def _process_paragraph(self, paragraph) -> Dict[str, Any]:
        """
        Procesa un párrafo del documento.
        
        Args:
            paragraph: Párrafo de python-docx
            
        Returns:
            Dict con la información del párrafo
        """
        # Extraer texto y propiedades
        text = paragraph.text.strip()
        if not text:
            return None
            
        # Detectar si es un encabezado por el estilo
        is_heading = False
        heading_level = 0
        style_name = "Normal"
        try:
            if paragraph.style and paragraph.style.name:
                style_name = paragraph.style.name
                is_heading = paragraph.style.name.lower().startswith('heading') or 'título' in paragraph.style.name.lower()
                if is_heading:
                    # Extraer nivel del nombre del estilo (e.g., 'Heading 1' -> 1)
                    try:
                        level_match = re.search(r'\d+', paragraph.style.name)
                        if level_match:
                            heading_level = int(level_match.group())
                        else:
                            heading_level = 1
                    except (ValueError, IndexError):
                        heading_level = 1
                        
                # Verificar otros estilos que podrían indicar encabezados
                if not is_heading and (
                    'title' in paragraph.style.name.lower() or 
                    'subtitle' in paragraph.style.name.lower() or
                    'caption' in paragraph.style.name.lower()
                ):
                    is_heading = True
                    heading_level = 1 if 'title' in paragraph.style.name.lower() else 2
                    
        except Exception as e:
            logger.debug(f"Error al procesar estilo: {str(e)}")
        
        # Detectar alineación
        alignment = None
        is_centered = False
        try:
            if paragraph.alignment:
                alignment = str(paragraph.alignment)
                is_centered = paragraph.alignment == docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            logger.debug(f"Error al procesar alineación: {str(e)}")
        
        # Analizar formato de los runs
        is_bold = False
        is_italic = False
        is_caps = False
        fonts = []
        try:
            runs_with_text = [run for run in paragraph.runs if run.text.strip()]
            if runs_with_text:
                is_bold = all(run.bold for run in runs_with_text)
                is_italic = all(run.italic for run in runs_with_text)
                # Verificar si todo está en mayúsculas
                is_caps = text.isupper() and len(text) > 3
                # Recopilar información de fuentes
                for run in runs_with_text:
                    if run.font.name:
                        fonts.append(run.font.name)
        except Exception as e:
            logger.debug(f"Error al procesar formato de texto: {str(e)}")
        
        # Verificaciones adicionales para detectar encabezados
        if not is_heading:
            # Los textos en mayúsculas cortos suelen ser títulos
            if is_caps and len(text) < 100:
                is_heading = True
                heading_level = 1
            # Textos en negrita cortos también pueden ser títulos
            elif is_bold and len(text) < 80:
                is_heading = True
                heading_level = 2
            # Textos centrados cortos son candidatos a títulos
            elif is_centered and len(text) < 80:
                is_heading = True
                heading_level = 2
            # Textos que terminan con puntos suspensivos pueden ser títulos
            elif text.endswith('...') and len(text) < 80:
                is_heading = True
                heading_level = 2
        
        # Verificar por patrones comunes de títulos en el texto
        heading_patterns = [
            r'^(La|El|Los|Las|Un|Una) [A-ZÁ-Ú][a-zá-ú]+(\.{3})?$', # La Permanencia...
            r'^[A-ZÁ-Ú]{4,}$',  # INTRODUCCIÓN
            r'^Capítulo \d+',    # Capítulo 1
            r'^[A-ZÁ-Ú][a-zá-ú]+(:|\.{3})' # Naturaleza divina:
        ]
        
        if not is_heading and any(re.match(pattern, text) for pattern in heading_patterns):
            is_heading = True
            heading_level = 2
        
        return {
            'text': text,
            'is_heading': is_heading,
            'heading_level': heading_level,
            'is_bold': is_bold,
            'is_italic': is_italic,
            'is_caps': is_caps,
            'is_centered': is_centered,
            'style': style_name,
            'alignment': alignment,
            'fonts': ','.join(fonts) if fonts else None
        }
    
    def load(self) -> Dict[str, Any]:
        """
        Carga y procesa el archivo DOCX.
        
        Returns:
            Dict[str, Any]: Diccionario con 'blocks' y 'document_metadata'
        """
        logger.info(f"Iniciando carga de archivo DOCX: {self.file_path}")
        
        try:
            # Abre el documento
            doc = docx.Document(self.file_path)
            
            # Obtiene información de fuente y contexto
            fuente, contexto = self.get_source_info()
            logger.debug(f"Fuente: {fuente}, Contexto: {contexto}")
            
            # Extrae metadatos del documento
            title = None
            author = None
            creation_date = None
            modification_date = None
            core_properties_dict = {}
            
            try:
                # Extraer propiedades del documento
                if doc.core_properties.title:
                    title = doc.core_properties.title
                if doc.core_properties.author:
                    author = doc.core_properties.author
                if doc.core_properties.created:
                    creation_date = doc.core_properties.created
                if doc.core_properties.modified:
                    modification_date = doc.core_properties.modified
                
                # Crear diccionario con todas las propiedades disponibles
                core_properties_dict = {
                    'title': doc.core_properties.title,
                    'author': doc.core_properties.author,
                    'subject': doc.core_properties.subject,
                    'keywords': doc.core_properties.keywords,
                    'comments': doc.core_properties.comments,
                    'category': doc.core_properties.category,
                    'created': doc.core_properties.created.isoformat() if doc.core_properties.created else None,
                    'modified': doc.core_properties.modified.isoformat() if doc.core_properties.modified else None,
                    'last_modified_by': doc.core_properties.last_modified_by,
                    'revision': doc.core_properties.revision,
                    'version': doc.core_properties.version
                }
                # Filtrar valores None
                core_properties_dict = {k: v for k, v in core_properties_dict.items() if v is not None}
                
            except Exception as e:
                logger.warning(f"Error al extraer propiedades del documento: {str(e)}")
            
            # Extrae fecha de las propiedades del documento o del nombre del archivo
            fecha = self._extract_date_from_core_properties(doc) or self._extract_date_from_filename()
            logger.debug(f"Fecha extraída: {fecha}")
            
            # Calcula hash del archivo
            file_hash = _calculate_sha256(self.file_path)
            
            # Crear metadatos del documento
            document_metadata = {
                "nombre_archivo": self.file_path.name,
                "ruta_archivo_original": str(self.file_path.resolve()),
                "extension_archivo": self.file_path.suffix,
                "titulo_documento": title if title else self.file_path.stem,
                "autor_documento": author,
                "fecha_publicacion_documento": fecha,
                "fecha_creacion_documento": creation_date.isoformat() if creation_date else None,
                "fecha_modificacion_documento": modification_date.isoformat() if modification_date else None,
                "hash_documento_original": file_hash,
                "original_fuente": fuente,
                "original_contexto": contexto,
                "metadatos_adicionales_fuente": core_properties_dict
            }
            
            # Procesa el contenido y extrae bloques
            blocks = []
            logger.debug(f"Total paragraphs found in DOCX: {len(doc.paragraphs)}")
            
            for i, paragraph in enumerate(doc.paragraphs):
                logger.debug(f"Processing paragraph {i+1}/{len(doc.paragraphs)} - Raw text: '{paragraph.text[:100]}...'" ) # Log first 100 chars
                metadata = self._process_paragraph(paragraph)
                if metadata:
                    logger.debug(f"Paragraph {i+1} processed. Metadata: {metadata}")
                    blocks.append(metadata)
                else:
                    logger.debug(f"Paragraph {i+1} skipped (empty or no metadata after processing).")
            
            logger.debug(f"Final blocks content: {len(blocks)} blocks extracted")
            
            # Devolver la estructura esperada por el BaseLoader
            return {
                'blocks': blocks,
                'document_metadata': document_metadata
            }
            
        except PackageNotFoundError:
            error_msg = f"Error al abrir archivo: {self.file_path} (no es un archivo DOCX válido o está corrupto)"
            logger.error(error_msg)
            return {
                'blocks': [],
                'document_metadata': {
                    "nombre_archivo": self.file_path.name,
                    "ruta_archivo_original": str(self.file_path.resolve()),
                    "extension_archivo": self.file_path.suffix,
                    "titulo_documento": self.file_path.stem,
                    "hash_documento_original": None,
                    "error": error_msg
                }
            }
        except Exception as e:
            error_msg = f"Error al procesar documento DOCX: {str(e)}"
            logger.error(error_msg)
            return {
                'blocks': [],
                'document_metadata': {
                    "nombre_archivo": self.file_path.name,
                    "ruta_archivo_original": str(self.file_path.resolve()),
                    "extension_archivo": self.file_path.suffix,
                    "titulo_documento": self.file_path.stem,
                    "hash_documento_original": None,
                    "error": error_msg
                }
            }