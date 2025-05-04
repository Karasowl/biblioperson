import os
from pathlib import Path
import subprocess
import json
import re
from datetime import datetime

# Importaciones para conversores de diferentes formatos
try:
    import docx  # Para archivos .docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd  # Para archivos CSV y Excel
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    import PyPDF2  # Para archivos PDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import textract  # Para otros formatos si está disponible
    TEXTRACT_AVAILABLE = True
except ImportError:
    TEXTRACT_AVAILABLE = False

# --- Funciones para convertir diferentes formatos a texto/Markdown ---

def convert_docx_to_markdown(docx_path, md_path, log_callback=None):
    """Convierte DOCX a Markdown usando Pandoc."""
    try:
        # Comando para ejecutar Pandoc
        # -f docx: formato de entrada
        # -t markdown: formato de salida (GitHub Flavored Markdown es una buena opción)
        # --wrap=none: Evitar que Pandoc añada saltos de línea artificiales
        # -o md_path: archivo de salida
        # docx_path: archivo de entrada
        command = [
            'pandoc', '-f', 'docx', '-t', 'markdown',
            '--wrap=none', '-o', str(md_path), str(docx_path)
        ]
        if log_callback:
            log_callback(f"  -> Running Pandoc: {' '.join(command)}")

        # Ejecutar Pandoc
        result = subprocess.run(command, capture_output=True, text=True, check=False)

        # Verificar si Pandoc se ejecutó correctamente
        if result.returncode != 0:
            error_msg = f"ERROR: Pandoc falló al convertir {docx_path.name}. Código: {result.returncode}\nStderr: {result.stderr}"
            if log_callback: log_callback(f"  -> {error_msg}")
            # Intentar dar una pista si Pandoc no está instalado
            if "command not found" in result.stderr.lower() or "no se reconoce" in result.stderr.lower() or result.returncode == 127:
                 error_msg += "\n(Asegúrate de que Pandoc esté instalado y en el PATH del sistema. Visita https://pandoc.org/installing.html)"
            return False, error_msg

        if log_callback:
            log_callback(f"  -> Pandoc conversion successful: {md_path.name}")
        return True, f"Converted {docx_path.name} to Markdown successfully."

    except FileNotFoundError:
         # Error si el ejecutable de pandoc no se encuentra
         error_msg = f"ERROR: No se encontró el comando 'pandoc'. Asegúrate de que Pandoc esté instalado y en el PATH del sistema. Visita https://pandoc.org/installing.html"
         if log_callback: log_callback(f"  -> {error_msg}")
         return False, error_msg
    except Exception as e:
        error_msg = f"ERROR: Excepción inesperada durante la conversión con Pandoc de {docx_path.name}: {str(e)}"
        if log_callback: log_callback(f"  -> {error_msg}")
        return False, error_msg

def convert_docx_to_text(file_path):
    if not DOCX_AVAILABLE:
        return "ERROR: La librería 'python-docx' no está instalada."
    try:
        doc = docx.Document(file_path)
        full_text_parts = []
        # Importar QName aquí para evitar dependencia global si docx no está instalado
        from docx.oxml.ns import qn

        for para in doc.paragraphs:
            paragraph_text = ""
            # Iterar sobre los elementos dentro del párrafo (runs, breaks, etc.)
            for element in para._element.iterchildren():
                # Si es un elemento de texto (run)
                if element.tag == qn('w:r'):
                    # Dentro de un run, buscar texto (w:t) y saltos (w:br)
                    for run_element in element.iterchildren():
                        if run_element.tag == qn('w:t') and run_element.text:
                            paragraph_text += run_element.text
                        elif run_element.tag == qn('w:br'): # Salto de línea explícito (Shift+Enter)
                            paragraph_text += '\n'
                # Si es un salto de línea explícito directamente bajo el párrafo (menos común)
                elif element.tag == qn('w:br'):
                     paragraph_text += '\n'

            full_text_parts.append(paragraph_text)

        # Unir los párrafos con doble salto de línea
        return '\n\n'.join(full_text_parts)
    except Exception as e:
        # Capturar específicamente el ImportError si qn no se pudo importar
        if isinstance(e, ImportError):
             return "ERROR: La librería 'python-docx' no está instalada o completa."
        return f"ERROR: No se pudo procesar el archivo DOCX {file_path.name}: {str(e)}"

def convert_doc_to_text(file_path):
    if not TEXTRACT_AVAILABLE:
        return f"ERROR: Formato .doc no soportado. Instala 'textract' o convierte a .docx."
    try:
        text = textract.process(str(file_path)).decode('utf-8')
        return text
    except Exception as e:
        return f"ERROR: No se pudo procesar el archivo DOC {file_path.name} con textract: {str(e)}"

def convert_excel_to_text(file_path):
    if not PANDAS_AVAILABLE:
        return "ERROR: La librería 'pandas' (y 'openpyxl') no está instalada para leer Excel."
    try:
        xls = pd.ExcelFile(file_path)
        full_text = []
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name, header=None)
            sheet_text = df.applymap(str).agg(' '.join, axis=1).str.cat(sep='\n')
            full_text.append(f"--- Hoja: {sheet_name} ---\n{sheet_text}")
        return '\n\n'.join(full_text)
    except Exception as e:
        return f"ERROR: No se pudo procesar el archivo Excel {file_path.name}: {str(e)}"

def convert_csv_to_text(file_path):
    if not PANDAS_AVAILABLE:
        return "ERROR: La librería 'pandas' no está instalada para leer CSV."
    try:
        df = pd.read_csv(file_path, header=None)
        text = df.applymap(str).agg(' '.join, axis=1).str.cat(sep='\n')
        return text
    except Exception as e:
        return f"ERROR: No se pudo procesar el archivo CSV {file_path.name}: {str(e)}"

def convert_pdf_to_text(file_path):
    if not PDF_AVAILABLE:
        return "ERROR: La librería 'PyPDF2' no está instalada para leer PDF."
    try:
        text = ""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                extracted = page.extract_text()
                if extracted:
                    text += extracted
                    text += "\n--- Page Break ---\n"
        return text.strip()
    except Exception as e:
        return f"ERROR: No se pudo procesar el archivo PDF {file_path.name}: {str(e)}"

def convert_any_to_text(file_path):
    if not TEXTRACT_AVAILABLE:
        return "ERROR: La librería 'textract' no está instalada para procesar formatos adicionales."
    try:
        text = textract.process(str(file_path)).decode('utf-8')
        return text
    except Exception as e:
        return f"ERROR: No se pudo procesar el archivo {file_path.name} con textract: {str(e)}"

def extract_docx_metadata(docx_path):
    """Extrae metadatos de un documento DOCX, enfocándose exclusivamente en la fecha de creación original."""
    if not DOCX_AVAILABLE:
        return {"error": "python-docx no está disponible"}
    
    try:
        doc = docx.Document(docx_path)
        core_props = doc.core_properties
        
        # Extraer SOLO la fecha de creación original
        created = core_props.created
        
        # Convertir a formato ISO si existe
        created_str = created.strftime('%Y-%m-%d') if created else ""
        
        # Otros metadatos útiles
        author = core_props.author or ""
        title = core_props.title or ""
        
        return {
            "fecha_creacion": created_str,  # Solo fecha de creación
            "autor": author,
            "titulo": title
        }
    except Exception as e:
        return {"error": f"Error extrayendo metadatos: {str(e)}"} 